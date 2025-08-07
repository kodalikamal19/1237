"""
Enhanced Document Processor with improved text extraction, intelligent chunking,
and optimized handling for insurance, legal, HR, and compliance documents.
"""

import os
import io
import re
import gc
import tempfile
from typing import List, Dict, Any, Tuple, Optional
import requests
import pypdf
import logging
from dataclasses import dataclass
from src.utils.memory_manager import MemoryManager

# Import OCR and document processing libraries
try:
    from pdf2image import convert_from_bytes
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import docx2txt
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import spacy
    # Load lightweight English model for better text processing
    try:
        nlp = spacy.load("en_core_web_sm")
        SPACY_AVAILABLE = True
    except OSError:
        SPACY_AVAILABLE = False
except ImportError:
    SPACY_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentSection:
    """Represents a structured document section"""
    title: str
    content: str
    section_type: str
    importance_score: float
    start_pos: int
    end_pos: int
    metadata: Dict[str, Any] = None

@dataclass
class ProcessedDocument:
    """Represents a fully processed document with structured sections"""
    raw_text: str
    cleaned_text: str
    sections: List[DocumentSection]
    document_type: str
    processing_method: str
    quality_score: float
    metadata: Dict[str, Any] = None

class EnhancedDocumentProcessor:
    """
    Enhanced document processor with intelligent text extraction, 
    domain-aware section identification, and optimized chunking strategies.
    """
    
    def __init__(self):
        self.memory_manager = MemoryManager()
        
        # Domain-specific section patterns
        self.section_patterns = {
            'insurance': {
                'coverage': [
                    r'coverage\s*details?',
                    r'what\s*is\s*covered',
                    r'benefits?\s*covered',
                    r'scope\s*of\s*coverage'
                ],
                'exclusions': [
                    r'exclusions?',
                    r'what\s*is\s*not\s*covered',
                    r'limitations?',
                    r'restrictions?'
                ],
                'conditions': [
                    r'terms?\s*and\s*conditions?',
                    r'policy\s*conditions?',
                    r'general\s*conditions?'
                ],
                'claims': [
                    r'claims?\s*procedure',
                    r'how\s*to\s*claim',
                    r'claims?\s*process'
                ],
                'definitions': [
                    r'definitions?',
                    r'meaning\s*of\s*terms',
                    r'glossary'
                ]
            },
            'legal': {
                'parties': [
                    r'parties?\s*to\s*the\s*agreement',
                    r'contracting\s*parties?'
                ],
                'obligations': [
                    r'obligations?',
                    r'duties\s*and\s*responsibilities',
                    r'shall\s*.*\s*party'
                ],
                'liability': [
                    r'liability',
                    r'indemnification',
                    r'damages?'
                ],
                'termination': [
                    r'termination',
                    r'breach',
                    r'default'
                ]
            },
            'hr': {
                'benefits': [
                    r'employee\s*benefits?',
                    r'compensation',
                    r'salary\s*structure'
                ],
                'policies': [
                    r'company\s*policies?',
                    r'hr\s*policies?',
                    r'workplace\s*policies?'
                ],
                'procedures': [
                    r'procedures?',
                    r'process',
                    r'guidelines?'
                ]
            }
        }
        
        # Text quality indicators
        self.quality_indicators = {
            'high_quality': [
                r'\d+\.\s+',  # Numbered lists
                r'[A-Z][a-z]+:',  # Section headers
                r'[Ss]ection\s+\d+',  # Section references
                r'\$\d+',  # Currency amounts
                r'\d+%',  # Percentages
                r'\d+\s+(days?|months?|years?)',  # Time periods
            ],
            'low_quality': [
                r'[^\w\s]{10,}',  # Too many special characters
                r'\s{5,}',  # Excessive whitespace
                r'[A-Z]{10,}',  # Too many capitals
                r'(.)\1{5,}',  # Repeated characters
            ]
        }
    
    @MemoryManager.cleanup_decorator
    def process_document_from_url(self, url: str, document_type: str = "insurance") -> ProcessedDocument:
        """
        Process document from URL with enhanced extraction and structuring.
        """
        try:
            # Download document
            logger.info(f"Downloading document from: {url}")
            document_bytes = self._download_document(url)
            
            # Detect document format
            doc_format = self._detect_document_format(document_bytes, url)
            logger.info(f"Detected document format: {doc_format}")
            
            # Extract text based on format
            if doc_format == 'pdf':
                raw_text, processing_method = self._extract_text_from_pdf(document_bytes)
            elif doc_format == 'docx':
                raw_text, processing_method = self._extract_text_from_docx(document_bytes)
            else:
                # Assume PDF as fallback
                raw_text, processing_method = self._extract_text_from_pdf(document_bytes)
            
            # Clean up document bytes immediately
            del document_bytes
            gc.collect()
            
            # Process the extracted text
            processed_doc = self._process_extracted_text(
                raw_text, document_type, processing_method
            )
            
            logger.info(f"Document processed successfully using {processing_method}")
            logger.info(f"Quality score: {processed_doc.quality_score:.2f}")
            logger.info(f"Sections identified: {len(processed_doc.sections)}")
            
            return processed_doc
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise
    
    def _download_document(self, url: str) -> bytes:
        """Download document with enhanced error handling and validation"""
        try:
            # Validate URL
            if not url or not isinstance(url, str):
                raise ValueError("Invalid URL provided")
            
            url = url.strip()
            
            # Enhanced headers
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'application/pdf,application/vnd.openxmlformats-officedocument.wordprocessingml.document,application/octet-stream,*/*',
                'Accept-Language': 'en-US,en;q=0.9',
                'Accept-Encoding': 'gzip, deflate, br',
                'Connection': 'keep-alive',
            }
            
            # Download with session
            session = requests.Session()
            session.headers.update(headers)
            
            response = session.get(url, stream=True, timeout=120)
            response.raise_for_status()
            
            # Check content type and size
            content_type = response.headers.get('content-type', '').lower()
            content_length = response.headers.get('content-length')
            
            if content_length:
                size_mb = int(content_length) / (1024 * 1024)
                if size_mb > 100:  # 100MB limit
                    raise ValueError(f"Document too large: {size_mb:.2f}MB")
            
            # Read in chunks
            document_data = io.BytesIO()
            total_size = 0
            max_size = 100 * 1024 * 1024  # 100MB
            
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    total_size += len(chunk)
                    if total_size > max_size:
                        document_data.close()
                        session.close()
                        raise ValueError("Document too large")
                    
                    document_data.write(chunk)
            
            document_bytes = document_data.getvalue()
            document_data.close()
            session.close()
            
            logger.info(f"Downloaded document: {len(document_bytes)} bytes")
            return document_bytes
            
        except Exception as e:
            raise Exception(f"Failed to download document: {str(e)}")
    
    def _detect_document_format(self, document_bytes: bytes, url: str = "") -> str:
        """Detect document format from bytes and URL"""
        
        # Check URL extension first
        if url:
            url_lower = url.lower()
            if '.pdf' in url_lower:
                return 'pdf'
            elif '.docx' in url_lower:
                return 'docx'
            elif '.doc' in url_lower:
                return 'doc'
        
        # Check file signature (magic numbers)
        if document_bytes.startswith(b'%PDF'):
            return 'pdf'
        elif document_bytes.startswith(b'PK\x03\x04') and b'word/' in document_bytes[:1000]:
            return 'docx'
        elif document_bytes.startswith(b'\xd0\xcf\x11\xe0'):
            return 'doc'
        
        # Default to PDF
        return 'pdf'
    
    def _extract_text_from_pdf(self, pdf_bytes: bytes) -> Tuple[str, str]:
        """Enhanced PDF text extraction with OCR fallback"""
        try:
            # Try standard text extraction first
            standard_text = self._extract_pdf_standard(pdf_bytes)
            
            if standard_text and len(standard_text.strip()) > 100:
                return standard_text, "PDF Standard Extraction"
            
            # Try OCR if standard extraction fails
            if OCR_AVAILABLE:
                logger.info("Standard extraction yielded little content, trying OCR...")
                ocr_text = self._extract_pdf_ocr(pdf_bytes)
                if ocr_text and len(ocr_text.strip()) > len(standard_text.strip()):
                    return ocr_text, "PDF OCR Extraction"
            
            if not standard_text or len(standard_text.strip()) < 50:
                raise Exception("No readable text found in PDF")
            
            return standard_text, "PDF Standard Extraction (Limited)"
            
        except Exception as e:
            raise Exception(f"PDF text extraction failed: {str(e)}")
    
    def _extract_pdf_standard(self, pdf_bytes: bytes) -> str:
        """Standard PDF text extraction using pypdf"""
        try:
            pdf_stream = io.BytesIO(pdf_bytes)
            reader = pypdf.PdfReader(pdf_stream)
            
            # Handle encryption
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except:
                    raise Exception("PDF is password-protected")
            
            # Extract text from all pages
            text_parts = []
            num_pages = min(len(reader.pages), 1000)  # Limit pages
            
            for page_num in range(num_pages):
                try:
                    page = reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text and page_text.strip():
                        cleaned_text = self._clean_extracted_text(page_text)
                        if cleaned_text:
                            text_parts.append(cleaned_text)
                    
                    # Memory management
                    if page_num % 20 == 0:
                        gc.collect()
                        
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num + 1}: {str(e)}")
                    continue
            
            pdf_stream.close()
            
            if not text_parts:
                return ""
            
            full_text = "\n\n".join(text_parts)
            return self._post_process_text(full_text)
            
        except Exception as e:
            logger.error(f"Standard PDF extraction failed: {str(e)}")
            return ""
    
    def _extract_pdf_ocr(self, pdf_bytes: bytes) -> str:
        """OCR-based PDF text extraction"""
        if not OCR_AVAILABLE:
            return ""
        
        try:
            # Convert PDF to images
            images = convert_from_bytes(
                pdf_bytes,
                dpi=200,
                first_page=1,
                last_page=50,  # Limit pages for OCR
                fmt='jpeg',
                thread_count=2
            )
            
            text_parts = []
            for i, image in enumerate(images):
                try:
                    # OCR with optimized settings
                    page_text = pytesseract.image_to_string(
                        image,
                        config='--psm 6 --oem 3 -l eng'
                    )
                    
                    if page_text and page_text.strip():
                        cleaned_text = self._clean_extracted_text(page_text)
                        if cleaned_text:
                            text_parts.append(cleaned_text)
                    
                    image.close()
                    
                    if i % 5 == 0:
                        gc.collect()
                        
                except Exception as e:
                    logger.warning(f"OCR failed for page {i+1}: {str(e)}")
                    continue
            
            # Clean up
            for img in images:
                try:
                    img.close()
                except:
                    pass
            del images
            gc.collect()
            
            if not text_parts:
                return ""
            
            full_text = "\n\n".join(text_parts)
            return self._post_process_ocr_text(full_text)
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            return ""
    
    def _extract_text_from_docx(self, docx_bytes: bytes) -> Tuple[str, str]:
        """Extract text from DOCX document"""
        if not DOCX_AVAILABLE:
            raise Exception("DOCX processing not available")
        
        try:
            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(docx_bytes)
                temp_file_path = temp_file.name
            
            try:
                # Try docx2txt first (simpler)
                text = docx2txt.process(temp_file_path)
                
                if text and len(text.strip()) > 50:
                    cleaned_text = self._clean_extracted_text(text)
                    return self._post_process_text(cleaned_text), "DOCX docx2txt"
                
                # Try python-docx for better structure
                doc = DocxDocument(temp_file_path)
                paragraphs = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text.strip())
                
                if paragraphs:
                    full_text = "\n\n".join(paragraphs)
                    cleaned_text = self._clean_extracted_text(full_text)
                    return self._post_process_text(cleaned_text), "DOCX python-docx"
                
                raise Exception("No text extracted from DOCX")
                
            finally:
                # Clean up temp file
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
                    
        except Exception as e:
            raise Exception(f"DOCX text extraction failed: {str(e)}")
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean extracted text with enhanced preprocessing"""
        if not text:
            return ""
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common PDF artifacts
        text = re.sub(r'[^\w\s.,;:!?()\[\]{}"\'/@#$%^&*+=<>~`|-]', '', text)
        
        # Fix common OCR errors
        ocr_corrections = {
            r'\|': 'I',
            r'\bO(?=\d)': '0',  # O before numbers
            r'(?<=\d)O\b': '0',  # O after numbers
            r'\brn\b': 'm',
            r'\bvv\b': 'w',
        }
        
        for pattern, replacement in ocr_corrections.items():
            text = re.sub(pattern, replacement, text)
        
        return text.strip()
    
    def _post_process_text(self, text: str) -> str:
        """Post-process text for better readability and structure"""
        if not text:
            return ""
        
        # Fix line breaks and spacing
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'([.!?])\s*\n\s*([A-Z])', r'\1\n\n\2', text)
        
        # Fix punctuation spacing
        text = re.sub(r'\s+([.,:;!?])', r'\1', text)
        text = re.sub(r'([.!?])\s*([a-z])', r'\1 \2', text)
        
        # Fix word boundaries
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        text = re.sub(r'(\d)([A-Za-z])', r'\1 \2', text)
        text = re.sub(r'([A-Za-z])(\d)', r'\1 \2', text)
        
        return text.strip()
    
    def _post_process_ocr_text(self, text: str) -> str:
        """Additional post-processing for OCR text"""
        text = self._post_process_text(text)
        
        # OCR-specific corrections
        text = re.sub(r'\b([A-Z])\s+([A-Z])\s+([A-Z])\b', r'\1\2\3', text)  # Fix spaced acronyms
        text = re.sub(r'(\d)\s+(\d)', r'\1\2', text)  # Fix spaced numbers
        
        return text
    
    def _process_extracted_text(self, raw_text: str, document_type: str, processing_method: str) -> ProcessedDocument:
        """Process extracted text into structured document"""
        
        # Clean the text
        cleaned_text = self._advanced_text_cleaning(raw_text)
        
        # Calculate quality score
        quality_score = self._calculate_text_quality(cleaned_text)
        
        # Identify sections
        sections = self._identify_document_sections(cleaned_text, document_type)
        
        # Create processed document
        processed_doc = ProcessedDocument(
            raw_text=raw_text,
            cleaned_text=cleaned_text,
            sections=sections,
            document_type=document_type,
            processing_method=processing_method,
            quality_score=quality_score,
            metadata={
                'original_length': len(raw_text),
                'cleaned_length': len(cleaned_text),
                'sections_count': len(sections),
                'avg_section_importance': sum(s.importance_score for s in sections) / len(sections) if sections else 0
            }
        )
        
        return processed_doc
    
    def _advanced_text_cleaning(self, text: str) -> str:
        """Advanced text cleaning with NLP if available"""
        
        # Basic cleaning
        cleaned = self._post_process_text(text)
        
        # Use spaCy for advanced cleaning if available
        if SPACY_AVAILABLE and len(cleaned) < 1000000:  # Limit for performance
            try:
                doc = nlp(cleaned[:100000])  # Process first 100k chars
                
                # Extract sentences for better structure
                sentences = []
                for sent in doc.sents:
                    sent_text = sent.text.strip()
                    if len(sent_text) > 10:  # Filter very short sentences
                        sentences.append(sent_text)
                
                if sentences:
                    cleaned = ' '.join(sentences)
                    
            except Exception as e:
                logger.warning(f"spaCy processing failed: {str(e)}")
        
        return cleaned
    
    def _calculate_text_quality(self, text: str) -> float:
        """Calculate text quality score"""
        if not text:
            return 0.0
        
        score = 0.5  # Base score
        
        # Check for quality indicators
        for pattern in self.quality_indicators['high_quality']:
            if re.search(pattern, text):
                score += 0.1
        
        # Penalize for low quality indicators
        for pattern in self.quality_indicators['low_quality']:
            if re.search(pattern, text):
                score -= 0.1
        
        # Length-based scoring
        if 1000 < len(text) < 100000:  # Good length range
            score += 0.2
        elif len(text) < 100:  # Too short
            score -= 0.3
        
        # Word diversity
        words = text.split()
        unique_words = set(words)
        if words and len(unique_words) / len(words) > 0.3:  # Good diversity
            score += 0.1
        
        return max(0.0, min(1.0, score))
    
    def _identify_document_sections(self, text: str, document_type: str) -> List[DocumentSection]:
        """Identify and extract document sections based on type"""
        
        sections = []
        patterns = self.section_patterns.get(document_type, {})
        
        if not patterns:
            # Generic section identification
            return self._identify_generic_sections(text)
        
        # Domain-specific section identification
        for section_type, section_patterns in patterns.items():
            for pattern in section_patterns:
                matches = list(re.finditer(pattern, text, re.IGNORECASE))
                
                for match in matches:
                    start_pos = match.start()
                    
                    # Find section content (next 2000 chars or until next section)
                    end_pos = min(start_pos + 2000, len(text))
                    
                    # Try to find natural end point
                    section_end = text.find('\n\n', start_pos + 100, end_pos)
                    if section_end != -1:
                        end_pos = section_end
                    
                    content = text[start_pos:end_pos].strip()
                    
                    if len(content) > 50:  # Minimum content length
                        importance_score = self._calculate_section_importance(content, section_type)
                        
                        section = DocumentSection(
                            title=match.group(0),
                            content=content,
                            section_type=section_type,
                            importance_score=importance_score,
                            start_pos=start_pos,
                            end_pos=end_pos,
                            metadata={'pattern_matched': pattern}
                        )
                        
                        sections.append(section)
        
        # Remove overlapping sections (keep higher importance)
        sections = self._remove_overlapping_sections(sections)
        
        # If no sections found, create generic sections
        if not sections:
            sections = self._identify_generic_sections(text)
        
        return sections
    
    def _identify_generic_sections(self, text: str) -> List[DocumentSection]:
        """Identify generic sections when domain-specific patterns fail"""
        
        sections = []
        
        # Split by double newlines (paragraph breaks)
        paragraphs = text.split('\n\n')
        
        current_section = ""
        section_start = 0
        
        for i, paragraph in enumerate(paragraphs):
            if len(current_section) + len(paragraph) < 1000:  # Target section size
                current_section += paragraph + "\n\n"
            else:
                if current_section.strip():
                    # Create section
                    section_end = section_start + len(current_section)
                    importance = self._calculate_section_importance(current_section, 'general')
                    
                    section = DocumentSection(
                        title=f"Section {len(sections) + 1}",
                        content=current_section.strip(),
                        section_type='general',
                        importance_score=importance,
                        start_pos=section_start,
                        end_pos=section_end,
                        metadata={'method': 'generic_chunking'}
                    )
                    
                    sections.append(section)
                    
                    section_start = section_end
                
                current_section = paragraph + "\n\n"
        
        # Add final section
        if current_section.strip():
            section_end = section_start + len(current_section)
            importance = self._calculate_section_importance(current_section, 'general')
            
            section = DocumentSection(
                title=f"Section {len(sections) + 1}",
                content=current_section.strip(),
                section_type='general',
                importance_score=importance,
                start_pos=section_start,
                end_pos=section_end,
                metadata={'method': 'generic_chunking'}
            )
            
            sections.append(section)
        
        return sections
    
    def _calculate_section_importance(self, content: str, section_type: str) -> float:
        """Calculate importance score for a section"""
        
        score = 0.3  # Base score
        content_lower = content.lower()
        
        # Type-specific importance
        type_weights = {
            'coverage': 0.9,
            'exclusions': 0.8,
            'conditions': 0.7,
            'claims': 0.8,
            'definitions': 0.6,
            'obligations': 0.8,
            'liability': 0.9,
            'benefits': 0.8,
            'general': 0.5
        }
        
        base_weight = type_weights.get(section_type, 0.5)
        score += base_weight * 0.4
        
        # Content-based scoring
        important_keywords = [
            'coverage', 'benefit', 'exclusion', 'condition', 'requirement',
            'shall', 'must', 'will', 'liability', 'claim', 'procedure',
            'amount', 'percent', 'days', 'months', 'years'
        ]
        
        keyword_count = sum(1 for keyword in important_keywords if keyword in content_lower)
        score += min(0.3, keyword_count * 0.05)
        
        # Numerical information bonus
        numbers = re.findall(r'\d+', content)
        if numbers:
            score += min(0.2, len(numbers) * 0.02)
        
        return min(1.0, score)
    
    def _remove_overlapping_sections(self, sections: List[DocumentSection]) -> List[DocumentSection]:
        """Remove overlapping sections, keeping those with higher importance"""
        
        if len(sections) <= 1:
            return sections
        
        # Sort by start position
        sections.sort(key=lambda x: x.start_pos)
        
        non_overlapping = []
        
        for section in sections:
            overlap_found = False
            
            for existing in non_overlapping:
                # Check for overlap
                if (section.start_pos < existing.end_pos and 
                    section.end_pos > existing.start_pos):
                    
                    # Keep section with higher importance
                    if section.importance_score > existing.importance_score:
                        non_overlapping.remove(existing)
                        non_overlapping.append(section)
                    
                    overlap_found = True
                    break
            
            if not overlap_found:
                non_overlapping.append(section)
        
        return non_overlapping